"""Pull text out of whatever the user pointed us at.

A "folder of corpus data" in practice means a pile of mixed formats: notes in
Markdown, scraped HTML, a few hundred PDFs, an exported JSONL of chat logs. Each
extractor yields `Record`s; nothing downstream needs to know where they came from.

Extractors never raise on a bad file. A corpus with three corrupt PDFs in it should
ingest the other 997 and tell you about the three.
"""

from __future__ import annotations

import csv
import io
import json
import sys
import unicodedata
import zipfile
from collections.abc import Iterator
from dataclasses import dataclass, field
from pathlib import Path

# Some corpora contain single CSV rows or JSON blobs of enormous size.
csv.field_size_limit(min(sys.maxsize, 2**31 - 1))


@dataclass(slots=True)
class Record:
    """One extracted unit: either prose or a conversation, never both."""

    source: str
    text: str | None = None
    messages: list[dict[str, str]] | None = None
    meta: dict = field(default_factory=dict)

    @property
    def is_chat(self) -> bool:
        return self.messages is not None

    def char_len(self) -> int:
        if self.messages is not None:
            return sum(len(m.get("content", "")) for m in self.messages)
        return len(self.text or "")


class SkipFile(Exception):
    """Raised by an extractor when a file has nothing usable in it."""


# ---------------------------------------------------------------------------
# extension routing
# ---------------------------------------------------------------------------

PROSE_EXT = {
    ".txt", ".text", ".md", ".markdown", ".mdx", ".rst", ".org",
    ".tex", ".srt", ".vtt", ".log", ".asc",
}

CODE_EXT = {
    ".py", ".pyi", ".js", ".jsx", ".ts", ".tsx", ".java", ".c", ".h", ".cpp",
    ".hpp", ".cc", ".cs", ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".kts",
    ".scala", ".sh", ".bash", ".zsh", ".fish", ".sql", ".r", ".lua", ".pl",
    ".vim", ".el", ".clj", ".ex", ".exs", ".erl", ".hs", ".ml", ".jl", ".dart",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".proto", ".graphql",
}

HTML_EXT = {".html", ".htm", ".xhtml", ".xml"}
JSONL_EXT = {".jsonl", ".ndjson"}
TABLE_EXT = {".csv", ".tsv"}

# Anything here is skipped silently rather than reported as an error — users point
# at real directories that happen to contain images and model weights.
BINARY_EXT = {
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".svg", ".webp", ".ico", ".tiff",
    ".mp3", ".mp4", ".wav", ".avi", ".mov", ".mkv", ".flac", ".ogg", ".webm",
    ".zip", ".tar", ".gz", ".bz2", ".xz", ".7z", ".rar", ".iso", ".dmg",
    ".exe", ".dll", ".so", ".dylib", ".bin", ".o", ".a", ".class", ".pyc",
    ".pt", ".pth", ".ckpt", ".safetensors", ".onnx", ".gguf", ".npy", ".npz",
    ".db", ".sqlite", ".sqlite3", ".parquet", ".arrow", ".feather",
    ".woff", ".woff2", ".ttf", ".otf", ".eot",
}

# Directories we never descend into.
SKIP_DIRS = {
    ".git", ".hg", ".svn", "node_modules", "__pycache__", ".venv", "venv",
    ".mypy_cache", ".pytest_cache", ".ruff_cache", ".tox", "dist", "build",
    ".idea", ".vscode", ".ipynb_checkpoints", ".cache", "site-packages",
}


def supported_extensions() -> set[str]:
    return (
        PROSE_EXT | CODE_EXT | HTML_EXT | JSONL_EXT | TABLE_EXT
        | {".json", ".pdf", ".docx", ".epub", ".ipynb"}
    )


# ---------------------------------------------------------------------------
# decoding
# ---------------------------------------------------------------------------


# chardet is excellent at statistically distinctive encodings (Shift-JIS, GB2312,
# KOI8-R) and unreliable at telling the Windows-125x family apart — on Western text
# with a handful of accents it reports ~0.1 confidence and picks the wrong codepage.
# So we trust it only when it is sure, and otherwise fall through to cp1252.
_CHARDET_MIN_CONFIDENCE = 0.8

_BOMS = (
    (b"\xef\xbb\xbf", "utf-8-sig"),
    (b"\xff\xfe\x00\x00", "utf-32-le"),
    (b"\x00\x00\xfe\xff", "utf-32-be"),
    (b"\xff\xfe", "utf-16-le"),
    (b"\xfe\xff", "utf-16-be"),
)


def decode_bytes(raw: bytes) -> str:
    """Best-effort decode. Never raises — real corpora are full of legacy encodings."""
    for bom, encoding in _BOMS:
        if raw.startswith(bom):
            try:
                return raw.decode(encoding)
            except UnicodeDecodeError:
                break

    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        pass

    try:
        import chardet

        guess = chardet.detect(raw[:200_000])
        encoding = guess.get("encoding")
        if encoding and (guess.get("confidence") or 0) >= _CHARDET_MIN_CONFIDENCE:
            try:
                return raw.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                pass
    except Exception:
        pass

    # cp1252 covers latin-1 plus the smart quotes and dashes that dominate legacy
    # Western text. It leaves five bytes undefined, so a strict failure here is a
    # real signal that the data is something else entirely.
    try:
        return raw.decode("cp1252")
    except UnicodeDecodeError:
        pass

    # latin-1 maps every byte, so this cannot fail.
    return raw.decode("latin-1")


def normalize(text: str) -> str:
    """Canonical form: NFC, unix line endings, no control characters."""
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Strip control chars but keep newline and tab.
    return "".join(c for c in text if c in "\n\t" or unicodedata.category(c)[0] != "C")


# ---------------------------------------------------------------------------
# structured records (json / jsonl / csv)
# ---------------------------------------------------------------------------

# Ordered by specificity: a record with both `messages` and `text` is a conversation
# that happens to carry a rendered copy, and the structured form wins.
_CHAT_LIST_KEYS = ("messages", "conversations", "conversation", "turns", "dialog", "dialogue")
_PROSE_KEYS = ("text", "content", "body", "document", "article", "raw", "page_content")
_PAIR_KEYS = (
    ("instruction", "output"),
    ("instruction", "response"),
    ("prompt", "completion"),
    ("prompt", "response"),
    ("question", "answer"),
    ("query", "response"),
    ("input", "output"),
)

_ROLE_KEYS = ("role", "from", "speaker")
_CONTENT_KEYS = ("content", "value", "text", "message")

_ROLE_ALIASES = {
    "human": "user",
    "prompter": "user",
    "user": "user",
    "gpt": "assistant",
    "assistant": "assistant",
    "bot": "assistant",
    "ai": "assistant",
    "model": "assistant",
    "system": "system",
}


def _normalize_turn(turn: object) -> dict[str, str] | None:
    if not isinstance(turn, dict):
        return None

    role = next((str(turn[k]) for k in _ROLE_KEYS if turn.get(k)), None)
    content = next((turn[k] for k in _CONTENT_KEYS if isinstance(turn.get(k), str)), None)
    if role is None or content is None:
        return None

    return {"role": _ROLE_ALIASES.get(role.strip().lower(), "user"), "content": content}


def record_from_mapping(obj: dict, source: str) -> Record | None:
    """Interpret one structured record, preferring conversation shapes over prose.

    Returns None when there is no usable text, which is common for metadata-only rows.
    """
    # 1. An explicit list of turns.
    for key in _CHAT_LIST_KEYS:
        value = obj.get(key)
        if isinstance(value, list) and value:
            turns = [t for t in (_normalize_turn(x) for x in value) if t]
            if turns:
                return Record(source=source, messages=turns)

    # 2. An instruction/response pair, optionally with a separate input field.
    for prompt_key, answer_key in _PAIR_KEYS:
        prompt, answer = obj.get(prompt_key), obj.get(answer_key)
        if isinstance(prompt, str) and isinstance(answer, str) and prompt.strip() and answer.strip():
            # Alpaca puts the operand in a sibling `input` field.
            extra = obj.get("input") if prompt_key != "input" else None
            if isinstance(extra, str) and extra.strip():
                prompt = f"{prompt}\n\n{extra}"

            messages = []
            system = obj.get("system")
            if isinstance(system, str) and system.strip():
                messages.append({"role": "system", "content": system})
            messages.append({"role": "user", "content": prompt})
            messages.append({"role": "assistant", "content": answer})
            return Record(source=source, messages=messages)

    # 3. Plain prose under a conventional key.
    for key in _PROSE_KEYS:
        value = obj.get(key)
        if isinstance(value, str) and value.strip():
            return Record(source=source, text=value)

    # 4. Last resort: a single string field, if the record has exactly one.
    strings = [v for v in obj.values() if isinstance(v, str) and len(v.strip()) > 32]
    if len(strings) == 1:
        return Record(source=source, text=strings[0])

    return None


# ---------------------------------------------------------------------------
# per-format extractors
# ---------------------------------------------------------------------------


def extract_prose(path: Path, source: str) -> Iterator[Record]:
    text = normalize(decode_bytes(path.read_bytes()))
    if text.strip():
        yield Record(source=source, text=text)


def extract_html(path: Path, source: str) -> Iterator[Record]:
    from selectolax.parser import HTMLParser

    tree = HTMLParser(decode_bytes(path.read_bytes()))
    for tag in tree.css("script, style, noscript, nav, footer, header, form"):
        tag.decompose()

    body = tree.body or tree.root
    if body is None:
        return
    text = normalize(body.text(separator="\n"))
    # Collapse the ragged blank-line runs that stripping tags leaves behind.
    text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
    if text.strip():
        yield Record(source=source, text=text)


def extract_pdf(path: Path, source: str) -> Iterator[Record]:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            continue  # one unparseable page shouldn't lose the document

    text = normalize("\n\n".join(pages))
    if not text.strip():
        raise SkipFile("no extractable text (likely a scanned image PDF)")
    yield Record(source=source, text=text, meta={"pages": len(reader.pages)})


def extract_docx(path: Path, source: str) -> Iterator[Record]:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))

    text = normalize("\n".join(parts))
    if text.strip():
        yield Record(source=source, text=text)


def extract_epub(path: Path, source: str) -> Iterator[Record]:
    """EPUBs are zipped XHTML; one Record per chapter keeps documents a sane size."""
    from selectolax.parser import HTMLParser

    with zipfile.ZipFile(path) as zf:
        names = [n for n in zf.namelist() if n.lower().endswith((".xhtml", ".html", ".htm"))]
        for name in sorted(names):
            try:
                tree = HTMLParser(decode_bytes(zf.read(name)))
            except Exception:
                continue
            for tag in tree.css("script, style"):
                tag.decompose()
            body = tree.body or tree.root
            if body is None:
                continue
            text = normalize(body.text(separator="\n"))
            text = "\n".join(line.strip() for line in text.split("\n") if line.strip())
            if len(text.strip()) > 128:
                yield Record(source=f"{source}#{name}", text=text)


def extract_ipynb(path: Path, source: str) -> Iterator[Record]:
    """Notebook source only. Outputs are mostly noise for language modelling."""
    nb = json.loads(decode_bytes(path.read_bytes()))
    parts = []
    for cell in nb.get("cells", []):
        src = cell.get("source", "")
        if isinstance(src, list):
            src = "".join(src)
        if not src.strip():
            continue
        if cell.get("cell_type") == "code":
            parts.append(f"```python\n{src}\n```")
        else:
            parts.append(src)

    text = normalize("\n\n".join(parts))
    if text.strip():
        yield Record(source=source, text=text)


def extract_jsonl(path: Path, source: str) -> Iterator[Record]:
    with path.open("rb") as f:
        for i, line in enumerate(f):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(decode_bytes(line))
            except json.JSONDecodeError:
                continue

            if isinstance(obj, str):
                yield Record(source=f"{source}:{i}", text=obj)
            elif isinstance(obj, dict):
                rec = record_from_mapping(obj, f"{source}:{i}")
                if rec:
                    yield rec


def extract_json(path: Path, source: str) -> Iterator[Record]:
    """Handles a top-level array, a wrapper dict around one, or a single record."""
    obj = json.loads(decode_bytes(path.read_bytes()))

    if isinstance(obj, dict):
        # Common wrapper shapes: {"data": [...]}, {"examples": [...]}.
        lists = [v for v in obj.values() if isinstance(v, list) and len(v) > 1]
        # Only unwrap if it isn't itself a conversation record.
        if lists and not any(k in obj for k in _CHAT_LIST_KEYS):
            obj = max(lists, key=len)

    items = obj if isinstance(obj, list) else [obj]
    for i, item in enumerate(items):
        if isinstance(item, str):
            yield Record(source=f"{source}:{i}", text=item)
        elif isinstance(item, dict):
            rec = record_from_mapping(item, f"{source}:{i}")
            if rec:
                yield rec


def extract_table(path: Path, source: str) -> Iterator[Record]:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    text = decode_bytes(path.read_bytes())
    reader = csv.DictReader(io.StringIO(text), delimiter=delimiter)

    if not reader.fieldnames:
        raise SkipFile("no header row")

    for i, row in enumerate(reader):
        clean = {k: v for k, v in row.items() if k and isinstance(v, str)}
        rec = record_from_mapping(clean, f"{source}:{i}")
        if rec:
            yield rec
            continue

        # No conventional key matched; fall back to the longest cell.
        longest = max(clean.values(), key=len, default="")
        if len(longest.strip()) > 64:
            yield Record(source=f"{source}:{i}", text=longest)


_EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".epub": extract_epub,
    ".ipynb": extract_ipynb,
    ".json": extract_json,
}


def extract(path: Path, source: str) -> Iterator[Record]:
    """Dispatch on extension. Raises SkipFile for unusable content."""
    ext = path.suffix.lower()

    if ext in _EXTRACTORS:
        yield from _EXTRACTORS[ext](path, source)
    elif ext in JSONL_EXT:
        yield from extract_jsonl(path, source)
    elif ext in TABLE_EXT:
        yield from extract_table(path, source)
    elif ext in HTML_EXT:
        yield from extract_html(path, source)
    elif ext in PROSE_EXT or ext in CODE_EXT:
        yield from extract_prose(path, source)
    else:
        raise SkipFile(f"unsupported extension '{ext}'")
